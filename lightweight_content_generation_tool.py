#!/usr/bin/env python
# coding: utf-8

import openai
import json,re
import os,gc,sys
from docx import Document

# ====== Configuration ======
openai.api_key = "your key here"  # Replace with your actual key
# MODEL = "gpt-4"
MODEL = "gpt-4.1"
# ===========================


def call_chatgpt(prompt):
    response = openai.ChatCompletion.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response['choices'][0]['message']['content']

def parse_json_response(response_text):
    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        print("⚠️ Error parsing JSON. Check the formatting of the API response.")
        return []

def export_to_word(columns, data, file_path, company_name):
    doc = Document()
    doc.add_heading("%s Content Strategy Plan"%company_name, 0)
    for idx, item in enumerate(data):
        doc.add_heading(item[columns[0]], level=1)
        doc.add_paragraph("Key Insights:"+'\n'+'\n'.join([x+';' for x in item[columns[1]]]))
        doc.add_paragraph("Objective:"+'\n'+item[columns[2]])
        doc.add_paragraph("Target Audience:"+'\n'+'\n'.join([x+';' for x in item[columns[3]].split(', ')]))
        doc.add_paragraph("Recommended Distribution Strategy:"+'\n'+'\n'.join([x+';' for x in item[columns[4]].split(', ')]))
        doc.add_paragraph("Priority Level (1=highest):"+'\n'+str(item[columns[5]]))
        doc.add_paragraph("Explanation:"+'\n'+str(item[columns[6]]))
    doc.save(file_path)
    print(f"✅ Word file saved at: {file_path}")

def main(company_name):
    PROMPT = """
Conduct thorough research on popular topics to identify emerging trends, analyze competitor strategies of %s company, and gather data-driven insights, focusing on 2024. 

Based on the research, generating engaging content ideas tailored to %s company according to below three main business objectives: 
1. Strengthening brand visibility
2. Attracting and hiring top talent
3. Reaching and engaging target clients

There are top ten different content formats, including Blog Posts, Videos, Social Media Content, Case Studies, Infographics, Webinars, Email Campaigns, E-books, User-Generated Content, Templates. 

Please select 10 content formats from the list and provide the following six items for each: (1) content format, (2) key insights in bullet points (the more the better), (3) relevant business objective (branding/recruitment/client acquisition), (4) target audience, (5) recommended distribution strategy, (6) priority level from 1 to 10 (1 = highest), and each item should be accompanied by a brief explanation (1–2 sentences) to clarify its relevance or strategic intent.

Please respond in a well-formatted JSON array.
"""%(company_name,company_name)

    print("⏳ Requesting response from ChatGPT...")
    raw_response = call_chatgpt(PROMPT)
    print("📥 Response received. Parsing JSON...")
    json_str = re.search(r"```json\n(.*)```", raw_response, re.DOTALL).group(1)
    content_list = json.loads(json_str)
    
    cols = list(content_list[0].keys())
    sorted_content = sorted(content_list, key=lambda x: x[cols[-2]])
    
    OUTPUT_PATH = r"./"  # Change to your target directory
    OUTPUT_FILE = "%s_content_plan.docx"%company_name

    output_full_path = os.path.join(OUTPUT_PATH, OUTPUT_FILE)
    export_to_word(cols, sorted_content, output_full_path, company_name)
    gc.collect()


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("❌ Please provide the company name as a command-line argument.")
        print(
            "Usage: python file.py CompanyName(capitalize first letter)\n"
            "Please make sure the company name is real and valid."
        )
        sys.exit(1)
    
    company_name = sys.argv[1]
    main(company_name)
